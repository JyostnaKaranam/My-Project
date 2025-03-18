import streamlit as st
st.set_page_config(layout="wide", initial_sidebar_state="collapsed")

import json
import torch

from transformers import BertTokenizer, BertModel

from PIL import Image

import pytesseract
from streamlit_option_menu import option_menu
import os
import pdfplumber
import plotly.graph_objects as go
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
import string
from dotenv import load_dotenv
import docx
from PIL import Image
from openai import OpenAI
import pygame
import speech_recognition as sr
from gtts import gTTS
import tempfile
import os
import random
import pandas as pd
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from transformers import T5Tokenizer, T5ForConditionalGeneration
import os
from transformers import BertModel, BertTokenizer
import torch
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
 
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
model = BertModel.from_pretrained('bert-base-uncased')
os.environ['TESSDATA_PREFIX'] = "C:\Program Files (x86)\Tesseract-OCR"

load_dotenv()
nltk.download('punkt')

session_state = st.session_state
if "user_index" not in st.session_state:
    st.session_state["user_index"] = 0
    st.session_state['current_question'] = None

if "verification_code" not in st.session_state:
    st.session_state["verification_code"] = None
if "email_verified" not in st.session_state:
    st.session_state.email_verified = False
if "email" not in st.session_state:
    st.session_state.email = ""


def signup(json_file_path="data.json"):
    st.title("Signup Page")
    name = st.text_input("Username:")
    email = st.text_input("Email:")
    sex = st.radio("Sex:", ("Male", "Female", "Other"))
    password = st.text_input("Password:", type="password")
    confirm_password = st.text_input("Confirm Password:", type="password")
    if st.button("Signup"):
        st.session_state.email = email
        if password == confirm_password:
            success, message = send_verification_email(email, session_state["verification_code"])
            if success:
                st.session_state.email_verified = False  
                st.session_state.show_verification = True
            else:
                st.error("Unable to send verification email at the moment. Please try again later.")
        else:
            st.error("Passwords do not match. Please try again.")
    if st.session_state.get("show_verification", False):
        verification = st.text_input("Enter Email Verification code")
        if st.button("Verify"):
            if verification.strip() == st.session_state.get("verification_code"):
                user = create_account(name, email, sex, password, json_file_path)
                session_state["logged_in"] = True
                session_state["user_info"] = user
                st.success("Account created and logged in successfully")
                session_state["verification_code"] = None
            else:
                st.error("Code doesn't match")
                st.write(st.session_state.get("verification_code"))


def check_login(username, password, json_file_path="data.json"):
    try:
        with open(json_file_path, "r") as json_file:
            data = json.load(json_file)

        for user in data["users"]:
            if user["name"] == username and user["password"] == password:
                session_state["logged_in"] = True
                session_state["user_info"] = user
                st.success("Login successful!")
                render_dashboard(user)
                return user

        st.error("Invalid credentials. Please try again.")
        return None
    except Exception as e:
        st.error(f"Error checking login: {e}")
        return None


def initialize_database(json_file_path="data.json"):
    try:
        if not os.path.exists(json_file_path):
            data = {"users": []}
            with open(json_file_path, "w") as json_file:
                json.dump(data, json_file)
    except Exception as e:
        print(f"Error initializing database: {e}")


def create_account(name, email, age, sex, password, json_file_path="data.json"):
    try:
        if not os.path.exists(json_file_path) or os.stat(json_file_path).st_size == 0:
            data = {"users": []}
        else:
            with open(json_file_path, "r") as json_file:
                data = json.load(json_file)

        user_info = {
            "name": name,
            "email": email,
            "sex": sex,
            "password": password,
            "resume": None,
            "job_description": None,
            "job_applied": None,
            'score': '0',
            "questions": None,
        }
        data["users"].append(user_info)


        with open(json_file_path, "w") as json_file:
            json.dump(data, json_file, indent=4)

        st.success("Account created successfully! You can now login.")
        return user_info
    except json.JSONDecodeError as e:
        st.error(f"Error decoding JSON: {e}")
        return None
    except Exception as e:
        st.error(f"Error creating account: {e}")
        return None


def login(json_file_path="data.json"):
    st.title("Login")
    username = st.text_input("Enter username:")
    password = st.text_input("Password:", type="password")

    login_button = st.button("Login")

    if login_button:
        user = check_login(username, password, json_file_path)
        if user is not None:
            session_state["logged_in"] = True
            session_state["user_info"] = user
        else:
            st.error("Invalid credentials. Please try again.")
            st.info("Forgot Password? Try resetting it.")
            if st.button("Reset Password"):
                show_forgot_password_page()


def get_user_info(email, json_file_path="data.json"):
    try:
        with open(json_file_path, "r") as json_file:
            data = json.load(json_file)
            for user in data["users"]:
                if user["email"] == email:
                    return user
        return None
    except Exception as e:
        st.error(f"Error getting user information: {e}")
        return None


def render_dashboard(user_info, json_file_path="data.json"):
    try:
        st.title(f"Welcome to the Dashboard, {user_info['name']}!")
        st.subheader("User Information:")
        st.write(f"Name: {user_info['name']}")
        st.write(f"Sex: {user_info['sex']}")
    except Exception as e:
        st.error(f"Error rendering dashboard: {e}")
def generate_reset_token():
    """Generate a random reset token"""
    return ''.join(random.choices(string.digits, k=6))


def save_reset_token(email, token):
    """Save reset token to CSV"""
    if not os.path.exists('reset_tokens.csv'):
        reset_df = pd.DataFrame(columns=['email', 'token', 'timestamp'])
    else:
        reset_df = pd.read_csv('reset_tokens.csv')

    reset_df['timestamp'] = pd.to_datetime(reset_df['timestamp'])

    reset_df = reset_df[reset_df['email'] != email]
    new_token = {
        'email': email,
        'token': token,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    reset_df = pd.concat([reset_df, pd.DataFrame([new_token])], ignore_index=True)
    reset_df.to_csv('reset_tokens.csv', index=False)
    return True


def send_verification_email(email, reset_token):
    """Send password reset email"""
    try:
        sender_email = "gowthamname907@gmail.com"
        sender_password = "fdnl bjao dduw sapy"

        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = email
        msg['Subject'] = 'Email Verification Request'

        reset_link = f"{reset_token}"

        body = f"""
        Hello,

        An Account Verification was requested for your account.

        Your verification code is: {reset_token}

        Please enter this code in the application to verify your account.

        If you did not request this, please ignore this email.
        """

        msg.attach(MIMEText(body, 'plain'))

        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
        server.quit()
        return True, "Email sent successfully"
    except Exception as e:
        return False, str(e)


def send_reset_email(email, reset_token):
    """Send password reset email"""
    try:
        sender_email = "gowthamname907@gmail.com"
        sender_password = "fdnl bjao dduw sapy"

        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = email
        msg['Subject'] = 'Password Reset Request'

        reset_link = f"{reset_token}"

        body = f"""
        Hello,

        A password reset was requested for your account.

        Your reset token is: {reset_token}

        Please enter this token in the application to reset your password.

        If you did not request this reset, please ignore this email.
        """

        msg.attach(MIMEText(body, 'plain'))

        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
        server.quit()
        return True, "Reset email sent successfully"
    except Exception as e:
        return False, str(e)


def verify_reset_token(email, token):
    """Verify if reset token is valid"""
    if not os.path.exists('reset_tokens.csv'):
        st.error("Tokens file not found")
        return False

    reset_df = pd.read_csv('reset_tokens.csv')
    token_row = reset_df[
        (reset_df['email'] == email) &
        (reset_df['token'] == token.strip())
        ]

    if token_row.empty:
        st.error("No matching data in tokens file")
        return False

    token_time = datetime.strptime(token_row.iloc[0]['timestamp'], '%Y-%m-%d %H:%M:%S')
    if (datetime.now() - token_time).total_seconds() > 24 * 3600:
        st.error("Token expired")
        return False

    return True


def update_password(email, new_password):
    """Update user's password"""
    with open('data.json', "r") as json_file:
        data = json.load(json_file)
        for user in data["users"]:
            if user["email"] == email:
                user['password'] = new_password
    with open('data.json', 'w') as json_file:
        json.dump(data, json_file, indent=4)
    return True, "Password Reset Successfully"


def show_forgot_password_page():
    """Display forgot password form"""
    st.markdown("""
        <style>
        .reset-status {
            margin-top: 1rem;
            padding: 1rem;
            border-radius: 5px;
        }
        </style>
    """, unsafe_allow_html=True)

    st.markdown("# 🔑 Reset Password")


    if 'reset_step' not in st.session_state:
        st.session_state['reset_step'] = 'email'

    if st.session_state['reset_step'] == 'email':
        email = st.text_input(" Enter your email address")

        if st.button("Send Reset Link", use_container_width=True):
            if not email:
                st.error("Please enter your email address")
                return

            with open('data.json', "r") as json_file:
                data = json.load(json_file)
                found = 0
                for user in data["users"]:
                    if user["email"] == email:
                        found = 1
                if not found:
                    st.error("Email address not found")
                    return

            reset_token = generate_reset_token()
            if save_reset_token(email, reset_token):

                success, message = send_reset_email(email, reset_token)
                if success:
                    st.success("Reset instructions sent to your email")
                    st.session_state['reset_email'] = email
                    st.session_state['reset_step'] = 'verify'
                    st.rerun()
                else:
                    st.error(f"Failed to send reset email: {message}")

    elif st.session_state['reset_step'] == 'verify':
        st.info(f" Reset token sent to {st.session_state['reset_email']}")
        token = st.text_input("🔑 Enter Reset Token")

        if st.button("Verify Token", use_container_width=True):
            if verify_reset_token(st.session_state['reset_email'], token):
                st.session_state['reset_step'] = 'reset'
                st.rerun()
            else:
                st.error("Invalid or expired token")

    elif st.session_state['reset_step'] == 'reset':
        new_password = st.text_input("🔒 Enter New Password", type="password")
        confirm_password = st.text_input("🔒 Confirm New Password", type="password")

        if st.button("Reset Password", use_container_width=True):
            if not new_password or not confirm_password:
                st.error("Please fill in all fields")
            elif new_password != confirm_password:
                st.error("Passwords do not match")
            else:
                success, message = update_password(st.session_state['reset_email'], new_password)
                if success:
                    st.success("Password reset successfully! Please login with your new password.")
                    del st.session_state['reset_step']
                    del st.session_state['reset_email']
                else:
                    st.error(f"Failed to reset password: {message}")

    st.markdown("</div>", unsafe_allow_html=True)

    if st.button("← Back to Login"):
        if 'reset_step' in st.session_state:
            del st.session_state['reset_step']
        if 'reset_email' in st.session_state:
            del st.session_state['reset_email']
        st.rerun()




def extract_text(file) -> str:
    if isinstance(file, str):
        file_extension = os.path.splitext(file)[1].lower()
    else:
        file_extension = os.path.splitext(file.name)[1].lower()


    if file_extension == '.pdf':
        if isinstance(file, str):
            with pdfplumber.open(file) as pdf:
                text = '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())
        else:
            with pdfplumber.open(file) as pdf:
                text = '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())


    elif file_extension == '.docx':
        if isinstance(file, str):
            doc = docx.Document(file)
        else:
            doc = docx.Document(file)
        text = '\n'.join([para.text for para in doc.paragraphs])


    elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
        if isinstance(file, str):
            image = Image.open(file)
        else:
            image = Image.open(file)
        text = pytesseract.image_to_string(image)

    else:
        if isinstance(file, str):
            with open(file, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        else:
            with file as f:
                text = f.read()

    return text
def extract_keywords_from_resume(resume_text):
    resume_text = resume_text.lower()
    resume_tokens = word_tokenize(resume_text)
    resume_tokens = [
        token for token in resume_tokens if token not in string.punctuation
    ]
    stop_words = set(stopwords.words("english"))
    resume_tokens = [token for token in resume_tokens if token not in stop_words]
    processed_resume_text = " ".join(resume_tokens)
    client = OpenAI(
        api_key=os.environ.get("OPENAI_API_KEY"),
    )
    prompt = f"Extract top most important skill keywords from the given resume text:\n{processed_resume_text}\nKeywords:"
    messages = [{"role": "system", "content": prompt}]
    response = client.chat.completions.create(
        messages=messages,
        model="gpt-3.5-turbo",
    )
    return response.choices[0].message.content.split(",")


def process_text(text):
    tokens = word_tokenize(text)
    tokens = [word.lower() for word in tokens if word.isalpha()]
    stop_words = set(stopwords.words('english'))
    tokens = [word for word in tokens if word not in stop_words]
    return tokens


def calculate_resume_score(job_description, resume):

    model = SentenceTransformer('all-MiniLM-L6-v2')

    job_desc_embedding = model.encode(job_description)
    resume_embedding = model.encode(resume)

    similarity_score = cosine_similarity([job_desc_embedding], [resume_embedding])[0][0] * 100
    return similarity_score



def generate_question(resume_text, job_description_text, candidate_name, previous_response=None,
                    previous_question=None):
    prompt = f"Resume Text: {resume_text}\nJob Description: {job_description_text}\nCandidate Name: {candidate_name}\n"
    if previous_response and previous_question:
        prompt += f"Previous Response: {previous_response}\nPrevious Question: {previous_question}\n"
    client = OpenAI(
        api_key=os.environ.get("OPENAI_API_KEY"),
    )
    messages = [
        {"role": "system", "content": "You are the interviewer."},
        {"role": "system",
        "content": "You are interviewing a candidate. Ask a question based on the resume and job description. If the candidate has already answered a question, you can ask a follow-up question based on their response."},
        {"role": "user", "content": prompt}
    ]
    response = client.chat.completions.create(
        messages=messages,
        model="gpt-3.5-turbo",
    )
    return response.choices[0].message.content


def generate_questions(resume_text, job_description_text, candidate_name, json_file_path="data.json"):

    if "questions_list" not in st.session_state:
        st.session_state.questions_list = []

        for _ in range(5):
            previous_response = None
            previous_question = None
            if st.session_state.questions_list:
                previous_response = st.session_state.questions_list[-1]["response"]
                previous_question = st.session_state.questions_list[-1]["question"]

            question = generate_question(summarize_t5(resume_text),
                                        job_description_text,
                                        candidate_name,
                                        previous_response,
                                        previous_question)
            st.session_state.questions_list.append({"question": question, "response": None})

    return st.session_state.questions_list


def interview_section(user_info, json_file_path="data.json"):
    st.title("Interview Questions")

    if user_info["resume"] is None or user_info["job_description"] is None:
        st.warning("Please upload your resume and apply for a job to generate interview questions.")
        return

    if "current_question_index" not in st.session_state:
        st.session_state.current_question_index = 0
    if "current_response" not in st.session_state:
        st.session_state.current_response = None
    if "question_number" not in st.session_state:
        st.session_state.question_number = 1


    questions = generate_questions(user_info["resume"],
                                user_info["job_description"],
                                user_info["name"])

    if st.session_state.current_question_index < len(questions):
        current_question = questions[st.session_state.current_question_index]["question"]
        st.markdown("### Current Question")
        st.markdown(current_question)

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Play Question"):
                try:
                    audio_file = text_to_speech(current_question)
                    play_audio(audio_file)
                except AssertionError:
                    st.error('No question found')

        with col2:
            input_method = st.radio("Choose input method:", ["Text", "Voice"])

        response = None
        if input_method == "Text":
            response = st.text_input("Your response:")
            if response:
                st.session_state.current_response = response
        else:
            voice_response = get_audio_input_with_buttons(current_question)
            if voice_response:
                st.write("Your response:", voice_response)
                st.session_state.current_response = voice_response
                response = voice_response

        col3, col4 = st.columns(2)
        with col3:
            if st.button("Next Question"):
                if st.session_state.current_response and len(st.session_state.current_response.strip()) > 0:
                    questions[st.session_state.current_question_index]["response"] = st.session_state.current_response

                    
                    with open(json_file_path, "r+") as json_file:
                        data = json.load(json_file)
                        user_index = next((i for i, user in enumerate(data["users"])
                                        if user["email"] == user_info["email"]), None)
                        if user_index is not None:
                            data["users"][user_index]["questions"] = questions
                            json_file.seek(0)
                            json.dump(data, json_file, indent=4)
                            json_file.truncate()

                    
                    st.session_state.current_question_index += 1
                    st.session_state.current_response = None
                    st.rerun()
                else:
                    st.warning("Please provide a response before moving to the next question.")
    else:
        if st.button("Finish Interview"):

            with open(json_file_path, "r+") as json_file:
                data = json.load(json_file)
                user_index = next((i for i, user in enumerate(data["users"])
                                if user["email"] == user_info["email"]), None)
                if user_index is not None:
                    data["users"][user_index]["questions"] = questions[:st.session_state.current_question_index + 1]
                    json_file.seek(0)
                    json.dump(data, json_file, indent=4)
                    json_file.truncate()

            st.success("Interview completed successfully!")
            del st.session_state.questions_list
            del st.session_state.current_question_index
            del st.session_state.current_response


def text_to_speech(text):
    """Convert text to speech and return the temporary file path"""
    tts = gTTS(text=text, lang='en')
    with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as fp:
        temp_file = fp.name
        tts.save(temp_file)
    return temp_file


def play_audio(file_path):
    """Play audio from the given file path"""
    pygame.mixer.init()
    pygame.mixer.music.load(file_path)
    pygame.mixer.music.play()
    while pygame.mixer.music.get_busy():
        pygame.time.Clock().tick(10)
    pygame.mixer.music.unload()
    os.unlink(file_path)


def get_audio_input_with_buttons(current_question):
    """Get audio input from microphone with start/stop buttons"""
    if 'recording' not in st.session_state:
        st.session_state.recording = False
    if 'current_question' not in st.session_state:
        st.session_state.current_question = current_question

    recognizer = sr.Recognizer()

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Start Recording"):
            st.session_state.recording = True

    with col2:
        if st.button("Stop Recording"):
            st.session_state.recording = False

    placeholder = st.empty()

    if st.session_state.recording:
        with sr.Microphone() as source:
            recognizer.adjust_for_ambient_noise(source, duration=0.5)
            placeholder.write("Recording... Press 'Stop Recording' when finished.")
            try:
                audio = recognizer.listen(source, timeout=None, phrase_time_limit=None)
                placeholder.write("Processing your response...")
                try:
                    text = recognizer.recognize_google(audio)
                    return text
                except sr.UnknownValueError:
                    placeholder.error("Could not understand the audio. Try again.")
                    return None
                except sr.RequestError as e:
                    placeholder.error(f"Could not request results; {e}")
                    return None
            except Exception as e:
                placeholder.error(f"Error recording audio: {e}")
                return None

    return None



def load_t5_models():
    model = T5ForConditionalGeneration.from_pretrained(
        "T5/t5_model.pt"
    )
    tokenizer = T5Tokenizer.from_pretrained(
        "T5/t5_tokenizer.pt"
    )
    return model, tokenizer

def summarize_t5(text):
    model, tokenizer = load_t5_models()
    max_length = len(text) // 2
    inputs = tokenizer.encode(
        "summarize: " + text, return_tensors="pt", max_length=512, truncation=True
    )
    summary_ids = model.generate(
        inputs.to(model.device),
        max_length=max_length,
        min_length=30,
        length_penalty=2.0,
        num_beams=4,
        early_stopping=True,
    )
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    return summary

def main(json_file_path="data.json"):
    with st.sidebar: 
        st.title("Web Based Selector-Applicant Simulation System")       
        page = option_menu(
            menu_title='',
            options=['Signup/Login', 'Upload Resume', 'Resume Analysis', 'Generate Questions', 'Evaluate Scores'],
            icons=['person-circle', 'upload', 'file-earmark-text', 'question-circle', 'check-circle'],
            menu_icon='chat-text-fill',
            default_index=0, 
            styles={
                "container": {"padding": "5!important", "background-color": 'black'},
                "icon": {"color": "white", "font-size": "23px"}, 
                "nav-link": {"color": "white", "font-size": "20px", "text-align": "left", "margin": "0px", "--hover-color": "blue"},
                "nav-link-selected": {"background-color": "#02ab21"},
            }
        )

    if page == "Signup/Login":
        st.title("Web-based Selector Applicant Simulation System")
        st.subheader("Signup/Login")
        login_or_signup = st.radio(
            "Select an option", ("Login", "Signup", "Forgot Password?"), key="login_signup"
        )
        if login_or_signup == "Login":
            login(json_file_path)
        elif login_or_signup == "Forgot Password?":
            show_forgot_password_page()
        else:

            token = generate_reset_token()
            if session_state["verification_code"] is None:
                session_state["verification_code"] = token
            signup(json_file_path)

    elif page == "Upload Resume":
        if session_state.get("logged_in"):
            uploaded_file = st.file_uploader("Choose a file", type=None)
            if uploaded_file is not None:
                resume_text = extract_text(uploaded_file)
                st.write("File name: ", uploaded_file.name)
                st.success("File uploaded successfully!")
                st.image(Image.open('Images/logo.png'), use_container_width=True)
                with open(json_file_path, "r+") as json_file:
                    data = json.load(json_file)
                    user_index = next((i for i, user in enumerate(data["users"]) if
                                    user["email"] == session_state["user_info"]["email"]), None)
                    if user_index is not None:
                        user_info = data["users"][user_index]
                        user_info["resume"] = resume_text
                        session_state["user_info"] = user_info
                        json_file.seek(0)
                        json.dump(data, json_file, indent=4)
                        json_file.truncate()
                    else:
                        st.error("User  not found.")

            st.subheader("Select a role to you want to apply for:")

            BASE_DIR = "Data/JobDesc/"
            job_description = st.selectbox(
                "Select a role",
                [
                    "-Select-",
                    "Backend Developer",
                    "Billing cum Logistics Manager",
                    "Data Scientist",
                    "Director of Engineering",
                    "Global Industry Content Manager",
                    "HTML Developer",
                    "IT Project Manager",
                    "Lead Technical Program Manager",
                    "Primary English Teacher",
                    "Revenue Reporting Data Analyst",
                    "Senior Product Manager",
                    "Senior Software Developer",
                    "Web Developer",
                    "Web_dev_job",
                ],
                key="job_description",
            )
            if job_description and job_description != "-Select-":
                file_path = os.path.join(BASE_DIR, f"{job_description}.docx")
                job_description_text = extract_text(file_path)
                st.subheader("Job Description:")
                st.write(job_description_text)
                if st.button("Submit"):
                    with open(json_file_path, "r+") as json_file:
                        data = json.load(json_file)
                        user_index = next((i for i, user in enumerate(data["users"]) if
                                        user["email"] == session_state["user_info"]["email"]), None)
                        if user_index is not None:
                            user_info = data["users"][user_index]
                            user_info["job_description"] = job_description_text
                            user_info["job_applied"] = job_description
                            session_state["user_info"] = user_info
                            json_file.seek(0)
                            json.dump(data, json_file, indent=4)
                            json_file.truncate()
                        else:
                            st.error("User  not found.")
                    st.success("submitted successfully!")

        else:
            st.warning("Please login/signup to view the dashboard.")


    elif page == "Resume Analysis":

        if session_state.get("logged_in"):

            st.title("Get Your Resume Analyzed and Compared")

            resume_text = session_state["user_info"]["resume"]

            resume_keywords = extract_keywords_from_resume(resume_text)

            resume_summary = summarize_t5(resume_text)

            st.subheader("Skills of the candidate:")

            for keyword in resume_keywords:
                st.write(f"- {keyword.strip()}")

            job_description_text = session_state["user_info"]["job_description"]

            resume_score = float(calculate_resume_score(job_description_text, resume_summary))

            with open(json_file_path, "r+") as json_file:
                data = json.load(json_file)
                user_index = next(
                    (i for i, user in enumerate(data["users"]) if user["email"] == session_state["user_info"]["email"]),
                    None)
                if user_index is not None:
                    user_info = data["users"][user_index]
                    user_info["score"] = float(resume_score) 
                    session_state["user_info"] = user_info
                    json_file.seek(0)
                    json.dump(data, json_file, indent=4)
                    json_file.truncate()
                else:
                    st.error("User  not found.")


            st.header("Resume Score")

            if resume_score >= 60:

                st.success(f"Congratulations! Your resume matches {resume_score:.2f}% with the job description.")

            elif resume_score >= 20:

                st.warning(
                    f"Your resume matches {resume_score:.2f}% with the job description. Consider improving it for better results.")

            else:

                st.error(
                    f"Your resume matches only {resume_score:.2f}% with the job description. Consider significant improvements.")


            percentage_score = resume_score / 100

            percentage_remainder = 1 - percentage_score

            fig = go.Figure(data=[go.Pie(labels=['Matched', 'Unmatched'],

                                        values=[percentage_score, percentage_remainder],

                                        hole=0.3,

                                        marker_colors=['rgba(0, 128, 0, 0.7)', 'rgba(255, 0, 0, 0.7'])])



            fig.update_layout(title_text="Resume Score")


            st.plotly_chart(fig)
            st.subheader("How does your resume compare with other candidates?")
            role = session_state["user_info"]["job_applied"]
            scores = [int(user["score"]) for user in data["users"] if user["job_applied"] == role]
            fig = go.Figure()
            fig.add_trace(go.Histogram(x=scores,
                                       histnorm='percent',
                                       marker_color='rgba(0, 0, 255, 0.7)',
                                       opacity=0.75))

            fig.update_layout(title_text=f'Distribution of Scores for Job Role: {role}',
                              xaxis_title='Resume Score',
                              yaxis_title='Percentage of Candidates',
                              bargap=0.05)

            st.plotly_chart(fig)

    elif page == "Generate Questions":
        if session_state.get("logged_in"):
            interview_section(session_state["user_info"])
        else:
            st.warning("Please login/signup to view the dashboard.")


    elif page == "Evaluate Scores":

        if session_state.get("logged_in"):

            user_info = session_state["user_info"]

            st.title("Evaluate Scores")

            st.subheader("Evaluate Scores")


            if user_info["questions"] is not None and len(user_info["questions"]) > 0:

                st.markdown("### Interview Questions and Responses")

                for question in user_info["questions"]:
                    st.markdown(f"Question: {question['question']}")

                    st.markdown(f"Response: {question['response']}")

                    st.markdown("---")

                st.subheader("Score")

                score = 0

                count = 0

                for question in user_info["questions"]:
                    count += 1

                    response = question["response"]

                    ques = question["question"]

                    score += calculate_resume_score(process_text(ques), process_text(response))


                if count > 0:

                    score = round(score / count, 2)

                else:

                    score = 0  

                st.write(f"Score: {score} %")

            else:

                st.warning("No questions available to evaluate.")


        else:

            st.warning("Please login/signup to view the dashboard.")

if __name__ == "__main__":
    initialize_database()
    nltk.download('punkt')
    nltk.download('stopwords')
    main()